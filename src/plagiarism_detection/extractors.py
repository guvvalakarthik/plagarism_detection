"""Bounded document text extraction."""

from __future__ import annotations

import io
from pathlib import Path

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 500_000


class DocumentExtractionError(ValueError):
    pass


def extract_document(filename: str, content: bytes) -> str:
    if not content:
        raise DocumentExtractionError("document is empty")
    if len(content) > MAX_UPLOAD_BYTES:
        raise DocumentExtractionError("document exceeds the 10 MB limit")
    extension = Path(filename).suffix.casefold()
    if extension == ".txt":
        text = _extract_text(content)
    elif extension == ".pdf":
        text = _extract_pdf(content)
    elif extension == ".docx":
        text = _extract_docx(content)
    else:
        raise DocumentExtractionError("only TXT, PDF, and DOCX are supported")
    text = text.replace("\x00", "").strip()
    if not text:
        raise DocumentExtractionError("document contains no extractable text")
    if len(text) > MAX_EXTRACTED_CHARACTERS:
        raise DocumentExtractionError("extracted text exceeds the character limit")
    return text


def _extract_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError("text encoding is not supported")


def _extract_pdf(content: bytes) -> str:
    if not content.startswith(b"%PDF"):
        raise DocumentExtractionError("file extension does not match PDF content")
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as error:
        raise DocumentExtractionError("PDF extraction failed") from error


def _extract_docx(content: bytes) -> str:
    if not content.startswith(b"PK"):
        raise DocumentExtractionError("file extension does not match DOCX content")
    try:
        from docx import Document

        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join([*paragraphs, *table_cells])
    except Exception as error:
        raise DocumentExtractionError("DOCX extraction failed") from error
