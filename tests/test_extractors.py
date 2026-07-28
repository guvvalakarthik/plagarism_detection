import io

import pytest
from docx import Document

from plagiarism_detection.extractors import (
    DocumentExtractionError,
    extract_document,
)


def test_extracts_utf8_text() -> None:
    assert extract_document("sample.txt", b"Evidence text") == "Evidence text"


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Paragraph evidence")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table evidence"
    buffer = io.BytesIO()
    document.save(buffer)

    extracted = extract_document("sample.docx", buffer.getvalue())

    assert "Paragraph evidence" in extracted
    assert "Table evidence" in extracted


@pytest.mark.parametrize(
    ("filename", "content"),
    [
        ("sample.exe", b"unsupported"),
        ("sample.pdf", b"not-a-pdf"),
        ("sample.docx", b"not-a-docx"),
        ("sample.txt", b""),
    ],
)
def test_rejects_invalid_documents(filename: str, content: bytes) -> None:
    with pytest.raises(DocumentExtractionError):
        extract_document(filename, content)
